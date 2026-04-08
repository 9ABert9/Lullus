"""
Export utilities for Lullus.

Exports content to Markdown, LaTeX, DOCX, Anki CSV, and PDF-ready formats.
"""

import csv
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


def export_to_markdown(
    content: str,
    title: str,
    metadata: Optional[Dict[str, Any]] = None,
) -> str:
    """Export content as Markdown with YAML frontmatter.

    Args:
        content: The main text content.
        title: Document title.
        metadata: Optional metadata dict for frontmatter.

    Returns:
        Complete Markdown string with frontmatter.
    """
    meta = metadata or {}
    date_str = meta.get("date", datetime.now().strftime("%Y-%m-%d"))
    author = meta.get("author", "Lullus")
    course = meta.get("course", "")
    tags = meta.get("tags", [])

    frontmatter_lines = [
        "---",
        f"title: \"{title}\"",
        f"author: \"{author}\"",
        f"date: {date_str}",
    ]
    if course:
        frontmatter_lines.append(f"course: \"{course}\"")
    if tags:
        frontmatter_lines.append(f"tags: [{', '.join(tags)}]")
    frontmatter_lines.append(f"generator: Lullus")
    frontmatter_lines.append("---")

    return "\n".join(frontmatter_lines) + "\n\n" + f"# {title}\n\n" + content + "\n"


def export_to_latex(
    content: str,
    title: str,
    metadata: Optional[Dict[str, Any]] = None,
) -> str:
    """Export content as a LaTeX document.

    Args:
        content: The main text content (Markdown-like).
        title: Document title.
        metadata: Optional metadata dict.

    Returns:
        Complete LaTeX document string.
    """
    meta = metadata or {}
    author = meta.get("author", "")
    date_str = meta.get("date", datetime.now().strftime("%Y-%m-%d"))

    # Basic Markdown-to-LaTeX conversion
    latex_content = _markdown_to_latex(content)

    return (
        "\\documentclass[12pt,a4paper]{article}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage{amsmath,amssymb}\n"
        "\\usepackage{hyperref}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage[margin=2.5cm]{geometry}\n"
        "\\usepackage{listings}\n"
        "\\usepackage{xcolor}\n"
        "\n"
        "\\lstset{\n"
        "  basicstyle=\\ttfamily\\small,\n"
        "  breaklines=true,\n"
        "  frame=single,\n"
        "  backgroundcolor=\\color{gray!10}\n"
        "}\n"
        "\n"
        f"\\title{{{_latex_escape(title)}}}\n"
        f"\\author{{{_latex_escape(author)}}}\n"
        f"\\date{{{_latex_escape(date_str)}}}\n"
        "\n"
        "\\begin{document}\n"
        "\\maketitle\n"
        "\n"
        f"{latex_content}\n"
        "\n"
        "\\end{document}\n"
    )


def _latex_escape(text: str) -> str:
    """Escape special LaTeX characters."""
    replacements = {
        "&": "\\&",
        "%": "\\%",
        "$": "\\$",
        "#": "\\#",
        "_": "\\_",
        "{": "\\{",
        "}": "\\}",
        "~": "\\textasciitilde{}",
        "^": "\\textasciicircum{}",
    }
    for char, escaped in replacements.items():
        text = text.replace(char, escaped)
    return text


def _markdown_to_latex(text: str) -> str:
    """Basic conversion of Markdown formatting to LaTeX."""
    lines = text.split("\n")
    result: List[str] = []
    in_code_block = False

    for line in lines:
        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                result.append("\\end{lstlisting}")
                in_code_block = False
            else:
                lang = line.strip().removeprefix("```").strip()
                if lang:
                    result.append(f"\\begin{{lstlisting}}[language={lang}]")
                else:
                    result.append("\\begin{lstlisting}")
                in_code_block = True
            continue

        if in_code_block:
            result.append(line)
            continue

        # Headings
        if line.startswith("#### "):
            result.append(f"\\paragraph{{{_latex_escape(line[5:].strip())}}}")
        elif line.startswith("### "):
            result.append(f"\\subsubsection{{{_latex_escape(line[4:].strip())}}}")
        elif line.startswith("## "):
            result.append(f"\\subsection{{{_latex_escape(line[3:].strip())}}}")
        elif line.startswith("# "):
            result.append(f"\\section{{{_latex_escape(line[2:].strip())}}}")
        elif line.startswith("- "):
            # Simple bullet conversion
            result.append(f"\\begin{{itemize}}\n\\item {_latex_escape(line[2:].strip())}\n\\end{{itemize}}")
        else:
            # Bold and italic
            converted = line
            # Bold: **text** -> \textbf{text}
            while "**" in converted:
                start = converted.index("**")
                rest = converted[start + 2:]
                if "**" in rest:
                    end = rest.index("**")
                    bold_text = rest[:end]
                    converted = (
                        converted[:start]
                        + f"\\textbf{{{_latex_escape(bold_text)}}}"
                        + rest[end + 2:]
                    )
                else:
                    break
            result.append(converted)

    if in_code_block:
        result.append("\\end{lstlisting}")

    return "\n".join(result)


def export_to_docx(
    content: str,
    title: str,
    metadata: Optional[Dict[str, Any]] = None,
    output_path: Optional[str] = None,
) -> str:
    """Export content as a DOCX file.

    Args:
        content: The main text content.
        title: Document title.
        metadata: Optional metadata dict.
        output_path: Path to save the DOCX file. If None, auto-generates.

    Returns:
        Path to the created DOCX file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    meta = metadata or {}
    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata subtitle
    author = meta.get("author", "")
    course = meta.get("course", "")
    date_str = meta.get("date", datetime.now().strftime("%Y-%m-%d"))
    if author or course:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_text = " | ".join(filter(None, [author, course, date_str]))
        run = subtitle.add_run(subtitle_text)
        run.font.size = Pt(10)
        run.font.italic = True

    doc.add_paragraph("")  # Spacer

    # Parse content by lines
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
            # Numbered list
            text = stripped.split(". ", 1)[1] if ". " in stripped else stripped
            doc.add_paragraph(text, style="List Number")
        else:
            para = doc.add_paragraph()
            # Handle bold text
            parts = stripped.split("**")
            for i, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    if i % 2 == 1:  # Odd indices are between ** markers
                        run.bold = True

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by Lullus")
    run.font.size = Pt(8)
    run.font.italic = True

    if not output_path:
        safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in title)
        output_path = f"{safe_title.strip().replace(' ', '_')}.docx"

    doc.save(output_path)
    logger.info("DOCX exported to: %s", output_path)
    return output_path


def export_to_anki_csv(
    flashcards: List[Dict[str, str]],
    output_path: str,
) -> str:
    """Export flashcards to Anki-compatible CSV.

    Args:
        flashcards: List of dicts with 'front' and 'back' keys.
        output_path: Path to save the CSV file.

    Returns:
        Path to the created CSV file.
    """
    path = Path(output_path)
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f, delimiter="\t")
        for card in flashcards:
            front = card.get("front", "")
            back = card.get("back", "")
            tags = card.get("tags", "")
            writer.writerow([front, back, tags])

    logger.info("Anki CSV exported to: %s (%d cards)", output_path, len(flashcards))
    return str(path)


def export_exercises_to_markdown(
    exercises: List[Any],
    title: str = "Exercises",
) -> str:
    """Export exercises to a formatted Markdown document.

    Args:
        exercises: List of Exercise objects with question, options, correct_answer, explanation, etc.
        title: Document title.

    Returns:
        Formatted Markdown string.
    """
    lines = [f"# {title}\n"]
    lines.append(f"*Generated by Lullus on {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n")

    for i, ex in enumerate(exercises, 1):
        question = getattr(ex, "question", str(ex))
        ex_type = getattr(ex, "exercise_type", "question")
        difficulty = getattr(ex, "difficulty", "medium")
        topic = getattr(ex, "topic", "")

        lines.append(f"## Question {i}")
        if topic:
            lines.append(f"**Topic:** {topic} | **Difficulty:** {difficulty} | **Type:** {ex_type}\n")
        lines.append(f"{question}\n")

        # Options for multiple choice
        options = getattr(ex, "options", None)
        if options:
            for opt in options:
                lines.append(f"- {opt}")
            lines.append("")

        lines.append("<details>")
        lines.append("<summary>Show Answer</summary>\n")

        correct = getattr(ex, "correct_answer", "")
        explanation = getattr(ex, "explanation", "")
        source = getattr(ex, "source_reference", "")

        if correct:
            lines.append(f"**Answer:** {correct}\n")
        if explanation:
            lines.append(f"**Explanation:** {explanation}\n")
        if source:
            lines.append(f"**Source:** {source}\n")

        lines.append("</details>\n")
        lines.append("---\n")

    return "\n".join(lines)
